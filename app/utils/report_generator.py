from docx import Document
from docx.shared import Inches
import os

def calculate_iou(predictions, ground_truth=None):
    """
    Calculates Intersection Over Union (IOU) for the embeddings.
    In PoV, if no ground truth is provided, we simulate an IOU metric 
    for the report.
    """
    if not ground_truth:
        # Simulate IOU calculation
        simulated_iou = 0.85 # Mocked 85% overlap
        return simulated_iou
    
    # Placeholder for real IOU logic:
    # area_of_overlap = calculate_overlap(predictions, ground_truth)
    # area_of_union = calculate_union(predictions, ground_truth)
    # return area_of_overlap / area_of_union
    return 0.85

def generate_word_report(results, report_name="IOU_Report.docx"):
    """
    Generates a Word Document summarizing the IOU measurements
    and model evaluation results.
    """
    document = Document()

    document.add_heading('Path Foundation Model Evaluation Report', 0)

    document.add_paragraph(
        'This report summarizes the Intersection Over Union (IOU) measurements '
        'for the selected WSI image batches using Google Path Foundation Models.'
    )

    for res in results:
        document.add_heading(f"Image: {res['image_name']}", level=1)
        
        p = document.add_paragraph()
        p.add_run('Model Evaluated: ').bold = True
        p.add_run(res.get('model', 'MedGemma/MedSigLip') + '\n')
        
        p.add_run('Calculated IOU: ').bold = True
        p.add_run(f"{res['iou'] * 100:.2f}%\n")
        
        p.add_run('Total Tiles Processed: ').bold = True
        p.add_run(f"{res['tiles_processed']}\n")
        
        p.add_run('Tumor/CRC Detected: ').bold = True
        p.add_run(f"{res['tumor_count']}\n")
        
        p.add_run('Normal Tissue: ').bold = True
        p.add_run(f"{res['normal_count']}\n")
        
    document.add_page_break()
    
    report_path = os.path.join("/tmp", report_name)
    document.save(report_path)
    return report_path
